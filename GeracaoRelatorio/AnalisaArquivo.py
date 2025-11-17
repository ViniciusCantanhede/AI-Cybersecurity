import pandas as pd
from docx import Document
from langchain_community.llms.ollama import Ollama
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate

llm = Ollama(model = "gemma3:1b")

#Criando parser para extrair apenas a resposta em texto
output_parser = StrOutputParser()

#função para gerar texto baseado nos dados do arquivo txt
def gera_resultado():
    df = pd.read_csv("log-web-server.txt", sep = r'\s+')
    resultados = []

    #criando pormpt
    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", "Você é um analista de segurança especializado em analisar logs de servidores web. Analise esses logs e forneça feedback em português brasileiro sobre possíveis anomalias ou tentativas de ataques e inclua recomendações de segurança para o servidor web"),
            ("user", "question: {question}")
        ]
    )

    #definindo cadeia de execução
    chain = prompt | llm | output_parser

    #criando documento word
    document = Document()
    document.add_heading('Relatório de Segurança a Partir do Log Web Server', 0)

    #itera sobre as linhas do dataframe
    for _, row in df.iterrows():
        #extrai os valores de cada linha com base nas colunas
        date, time, s_ip, cs_method, cs_uri_stem, cs_uri_query, s_port, cs_username, c_ip, cs_user_agent, cs_referer, sc_status, sc_substatus, sc_win32_status, time_taken = row
        #Cria prompt para o llm com base nos dados da execuçao do pipeline
        consulta_pipeline = (f"Data: {date}, Hora: {time}, IP de Origem: {s_ip}, Método: {cs_method}, URI: {cs_uri_stem}, " 
                             f"Porta: {s_port}, Usuário: {cs_username}, IP do Cliente: {c_ip}, Statu: {sc_status}, "
                             f"Tempo tomado: {time_taken}ms.")

        #gera o texto de resultado 
        response = chain.invoke({"question": consulta_pipeline})
        resultados.append(response)
        document.add_paragraph(response)
    #salva o documento word
    document.save('Relatorio_Seguranca_Log_Web_Server.docx')
    return resultados

resultados = gera_resultado()
print("Relatório de segurança gerado com sucesso: Relatorio_Seguranca_Log_Web_Server.docx")